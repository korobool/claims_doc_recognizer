"""Document + Page registry.

A *Document* is the top-level thing a user uploads. It can be:
    - a single image (png/jpeg/tiff/bmp)  -> 1-page Document
    - a PDF file                           -> N-page Document
    - a DOCX file                          -> 1-page Document (text-rendered)
    - a "multipage" grouping of N files    -> N-page Document

Each *Page* always has a raster image (`image_path`) the viewer can display,
and may additionally have a `text_layer` string lifted directly from the
source. When a page has a non-trivial text layer, downstream OCR can be
skipped and the text layer used verbatim — that's the "PDF with embedded OCR
layer" fast-path mentioned in the product spec.

Storage layout on disk:

    uploads/<doc_id>/source.<ext>    # original upload (pdf, docx, png, ...)
    uploads/<doc_id>/page_1.png      # rendered / copied pages
    uploads/<doc_id>/page_2.png
    ...

All registry state is in-memory and re-populated on ingest. File state survives
restarts but the in-memory index does not.
"""

from __future__ import annotations

import io
import os
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from PIL import Image, ImageDraw, ImageFont


UPLOADS_ROOT = Path(__file__).parent.parent.parent / "uploads"

# A page text layer must have at least this many non-whitespace characters for
# us to trust it and skip OCR. Scanned PDFs often contain a handful of stray
# characters per page (junk from the scanner's text layer) that we don't want
# to treat as real content.
TEXT_LAYER_MIN_CHARS = 40

SUPPORTED_IMAGE_MIMES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/tiff",
    "image/bmp",
    "image/webp",
}
SUPPORTED_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}


@dataclass
class Page:
    """One displayable unit belonging to a Document."""

    page_id: str
    doc_id: str
    index: int  # 1-based, human-friendly
    image_path: str
    text_layer: Optional[str] = None  # natively extracted text, if any
    width: int = 0
    height: int = 0

    @property
    def has_text_layer(self) -> bool:
        return bool(self.text_layer and len(self.text_layer.strip()) >= TEXT_LAYER_MIN_CHARS)

    def to_summary(self) -> Dict:
        return {
            "page_id": self.page_id,
            "doc_id": self.doc_id,
            "index": self.index,
            "has_text_layer": self.has_text_layer,
            "width": self.width,
            "height": self.height,
        }


@dataclass
class Document:
    doc_id: str
    filename: str
    source_kind: str  # "image" | "pdf" | "docx" | "multipage"
    pages: List[Page] = field(default_factory=list)
    source_file: Optional[str] = None  # path to original upload on disk

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def has_any_text_layer(self) -> bool:
        return any(p.has_text_layer for p in self.pages)

    def full_text_layer(self) -> str:
        """Concatenate all pages' text layers with page markers."""
        out: List[str] = []
        for p in self.pages:
            if p.has_text_layer:
                out.append(f"=== Page {p.index} ===\n{p.text_layer.strip()}")
        return "\n\n".join(out)

    def to_summary(self) -> Dict:
        return {
            "doc_id": self.doc_id,
            "filename": self.filename,
            "source_kind": self.source_kind,
            "page_count": self.page_count,
            "has_text_layer": self.has_any_text_layer,
            "pages": [p.to_summary() for p in self.pages],
        }


# -----------------------------------------------------------------------------
# Registry
# -----------------------------------------------------------------------------

_documents: Dict[str, Document] = {}
_pages: Dict[str, Page] = {}  # page_id -> Page (secondary index)


def register(doc: Document) -> Document:
    _documents[doc.doc_id] = doc
    for p in doc.pages:
        _pages[p.page_id] = p
    return doc


def get_document(doc_id: str) -> Optional[Document]:
    return _documents.get(doc_id)


def get_page(page_id: str) -> Optional[Page]:
    return _pages.get(page_id)


def list_documents() -> List[Document]:
    return list(_documents.values())


def delete_document(doc_id: str) -> bool:
    doc = _documents.pop(doc_id, None)
    if not doc:
        return False
    for p in doc.pages:
        _pages.pop(p.page_id, None)
    try:
        doc_dir = UPLOADS_ROOT / doc_id
        if doc_dir.exists():
            for f in doc_dir.iterdir():
                f.unlink()
            doc_dir.rmdir()
    except Exception as e:
        print(f"[DocumentService] failed to clean up {doc_id}: {e}")
    return True


def _new_doc_dir() -> Tuple[str, Path]:
    doc_id = str(uuid.uuid4())
    d = UPLOADS_ROOT / doc_id
    d.mkdir(parents=True, exist_ok=True)
    return doc_id, d


# -----------------------------------------------------------------------------
# Ingest: images
# -----------------------------------------------------------------------------


def ingest_image(filename: str, content: bytes) -> Document:
    """Save a single image as a 1-page Document."""
    doc_id, doc_dir = _new_doc_dir()
    ext = Path(filename).suffix.lower() or ".png"
    if ext not in SUPPORTED_IMAGE_EXTS:
        ext = ".png"
    source_path = doc_dir / f"source{ext}"
    with open(source_path, "wb") as f:
        f.write(content)

    # The page image is always PNG for uniform downstream handling.
    page_path = doc_dir / "page_1.png"
    try:
        with Image.open(io.BytesIO(content)) as im:
            im.load()
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGB")
            w, h = im.size
            im.save(page_path, "PNG")
    except Exception as e:
        # Fallback — write bytes as-is (may still be openable by Surya).
        page_path.write_bytes(content)
        w = h = 0
        print(f"[DocumentService] image normalise failed for {filename}: {e}")

    page = Page(
        page_id=str(uuid.uuid4()),
        doc_id=doc_id,
        index=1,
        image_path=str(page_path),
        text_layer=None,
        width=w,
        height=h,
    )
    doc = Document(
        doc_id=doc_id,
        filename=filename,
        source_kind="image",
        pages=[page],
        source_file=str(source_path),
    )
    return register(doc)


# -----------------------------------------------------------------------------
# Ingest: PDF
# -----------------------------------------------------------------------------


def ingest_pdf(filename: str, content: bytes) -> Document:
    """Render each PDF page to PNG and capture its text layer if present."""
    import pypdfium2 as pdfium

    doc_id, doc_dir = _new_doc_dir()
    source_path = doc_dir / "source.pdf"
    source_path.write_bytes(content)

    pages: List[Page] = []
    try:
        pdf = pdfium.PdfDocument(content)
    except Exception as e:
        raise ValueError(f"Unable to open PDF: {e}")

    for i in range(len(pdf)):
        page_obj = pdf[i]
        # Raster at ~200 DPI (scale 200/72 ≈ 2.78). Surya's detection is
        # sensitive to resolution — too low and text lines get merged.
        rendered = page_obj.render(scale=2.78).to_pil()
        if rendered.mode not in ("RGB", "RGBA"):
            rendered = rendered.convert("RGB")
        w, h = rendered.size
        page_path = doc_dir / f"page_{i + 1}.png"
        rendered.save(page_path, "PNG")

        # Extract text layer for this page. pypdfium2 returns the raw page
        # text — may be empty for scanned pages.
        text_page = page_obj.get_textpage()
        try:
            text = text_page.get_text_bounded() or ""
        finally:
            text_page.close()

        pages.append(
            Page(
                page_id=str(uuid.uuid4()),
                doc_id=doc_id,
                index=i + 1,
                image_path=str(page_path),
                text_layer=text if text.strip() else None,
                width=w,
                height=h,
            )
        )

    pdf.close()

    doc = Document(
        doc_id=doc_id,
        filename=filename,
        source_kind="pdf",
        pages=pages,
        source_file=str(source_path),
    )
    return register(doc)


# -----------------------------------------------------------------------------
# Ingest: DOCX
# -----------------------------------------------------------------------------


def ingest_docx(filename: str, content: bytes) -> Document:
    """Extract DOCX text and render a placeholder page image for the viewer.

    DOCX has no strict page breaks (layout depends on the rendering engine),
    so we treat the whole document as a single Page with a natively-available
    text layer. The page image is a best-effort render of the text so the
    viewer has something to show; the LLM and classifier work off the text.
    """
    import docx

    doc_id, doc_dir = _new_doc_dir()
    source_path = doc_dir / "source.docx"
    source_path.write_bytes(content)

    try:
        wd = docx.Document(io.BytesIO(content))
    except Exception as e:
        raise ValueError(f"Unable to open DOCX: {e}")

    paragraphs = [p.text for p in wd.paragraphs]
    # Tables → tab-separated rows
    for tbl in wd.tables:
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append("\t".join(cells))
    full_text = "\n".join(p for p in paragraphs if p is not None).strip()

    page_path = doc_dir / "page_1.png"
    _render_text_to_image(full_text or "(empty document)", page_path, title=filename)

    with Image.open(page_path) as im:
        w, h = im.size

    page = Page(
        page_id=str(uuid.uuid4()),
        doc_id=doc_id,
        index=1,
        image_path=str(page_path),
        text_layer=full_text if full_text else None,
        width=w,
        height=h,
    )
    doc = Document(
        doc_id=doc_id,
        filename=filename,
        source_kind="docx",
        pages=[page],
        source_file=str(source_path),
    )
    return register(doc)


def _render_text_to_image(text: str, out_path: Path, title: str = "") -> None:
    """Best-effort: draw the DOCX text onto a synthetic page so the viewer has
    something visual to show. This is NOT meant to be accurate typesetting."""
    W, H = 1200, 1700  # ~US Letter at 150 DPI
    im = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(im)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 28)
        body_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 18)
    except Exception:
        title_font = body_font = ImageFont.load_default()

    y = 60
    margin = 60
    if title:
        draw.text((margin, y), title, fill="black", font=title_font)
        y += 50
        draw.line([(margin, y), (W - margin, y)], fill="#888", width=2)
        y += 24

    # Soft line wrap at ~100 chars. The goal is legibility for the viewer,
    # not faithful Word rendering.
    max_chars = 95
    for paragraph in text.split("\n"):
        paragraph = paragraph.rstrip()
        if not paragraph:
            y += 22
            continue
        while paragraph:
            chunk, paragraph = paragraph[:max_chars], paragraph[max_chars:]
            draw.text((margin, y), chunk, fill="black", font=body_font)
            y += 24
            if y > H - 60:
                break
        if y > H - 60:
            draw.text(
                (margin, H - 50), "... (truncated for preview) ...", fill="#888", font=body_font
            )
            break

    im.save(out_path, "PNG")


# -----------------------------------------------------------------------------
# Ingest: multipage grouping
# -----------------------------------------------------------------------------


def ingest_multipage(filename: str, files: List[Tuple[str, bytes]]) -> Document:
    """Group arbitrary uploaded files into a single multi-page Document.

    Each input file may be an image, a PDF, or a DOCX. We flatten them into a
    single page sequence in the order provided. This lets the user staple
    together screenshots, or combine a cover page PDF with a scan, etc.
    """
    if not files:
        raise ValueError("No files provided for multipage document")

    doc_id, doc_dir = _new_doc_dir()
    pages: List[Page] = []
    next_index = 1

    # For multipage, we don't keep a single "source" file; instead we archive
    # each input under uploads/<doc_id>/part_<N>.<ext> for auditability.
    for i, (fname, content) in enumerate(files, start=1):
        part_ext = Path(fname).suffix.lower() or ".bin"
        part_path = doc_dir / f"part_{i}{part_ext}"
        part_path.write_bytes(content)

        mime = _guess_mime(fname, content)
        if mime == "application/pdf":
            pages.extend(_pages_from_pdf_bytes(doc_id, doc_dir, content, next_index))
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            pages.extend(_pages_from_docx_bytes(doc_id, doc_dir, content, fname, next_index))
        else:
            pages.extend(_pages_from_image_bytes(doc_id, doc_dir, content, next_index))
        next_index = len(pages) + 1

    if not pages:
        raise ValueError("Multipage upload produced no pages")

    doc = Document(
        doc_id=doc_id,
        filename=filename,
        source_kind="multipage",
        pages=pages,
        source_file=None,
    )
    return register(doc)


def _pages_from_image_bytes(doc_id: str, doc_dir: Path, content: bytes, start_index: int) -> List[Page]:
    page_path = doc_dir / f"page_{start_index}.png"
    try:
        with Image.open(io.BytesIO(content)) as im:
            im.load()
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGB")
            w, h = im.size
            im.save(page_path, "PNG")
    except Exception as e:
        page_path.write_bytes(content)
        w = h = 0
        print(f"[DocumentService] multipage image normalise failed: {e}")
    return [
        Page(
            page_id=str(uuid.uuid4()),
            doc_id=doc_id,
            index=start_index,
            image_path=str(page_path),
            text_layer=None,
            width=w,
            height=h,
        )
    ]


def _pages_from_pdf_bytes(doc_id: str, doc_dir: Path, content: bytes, start_index: int) -> List[Page]:
    import pypdfium2 as pdfium

    pages: List[Page] = []
    pdf = pdfium.PdfDocument(content)
    for i in range(len(pdf)):
        idx = start_index + i
        rendered = pdf[i].render(scale=2.78).to_pil()
        if rendered.mode not in ("RGB", "RGBA"):
            rendered = rendered.convert("RGB")
        w, h = rendered.size
        page_path = doc_dir / f"page_{idx}.png"
        rendered.save(page_path, "PNG")

        text_page = pdf[i].get_textpage()
        try:
            text = text_page.get_text_bounded() or ""
        finally:
            text_page.close()

        pages.append(
            Page(
                page_id=str(uuid.uuid4()),
                doc_id=doc_id,
                index=idx,
                image_path=str(page_path),
                text_layer=text if text.strip() else None,
                width=w,
                height=h,
            )
        )
    pdf.close()
    return pages


def _pages_from_docx_bytes(doc_id: str, doc_dir: Path, content: bytes, filename: str, start_index: int) -> List[Page]:
    import docx

    wd = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in wd.paragraphs]
    for tbl in wd.tables:
        for row in tbl.rows:
            paragraphs.append("\t".join(cell.text.strip() for cell in row.cells))
    full_text = "\n".join(p for p in paragraphs if p is not None).strip()
    page_path = doc_dir / f"page_{start_index}.png"
    _render_text_to_image(full_text or "(empty document)", page_path, title=filename)
    with Image.open(page_path) as im:
        w, h = im.size
    return [
        Page(
            page_id=str(uuid.uuid4()),
            doc_id=doc_id,
            index=start_index,
            image_path=str(page_path),
            text_layer=full_text if full_text else None,
            width=w,
            height=h,
        )
    ]


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------


def _guess_mime(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or content[:4] == b"%PDF":
        return "application/pdf"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext == ".doc":
        return "application/msword"
    if ext in SUPPORTED_IMAGE_EXTS:
        return f"image/{ext[1:].replace('jpg','jpeg')}"
    # Sniff by magic bytes as a last resort.
    if content[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if content[:2] == b"\xff\xd8":
        return "image/jpeg"
    return "application/octet-stream"


def classify_upload(filename: str, content: bytes) -> str:
    """Return one of 'image', 'pdf', 'docx', or 'unsupported' for the upload."""
    mime = _guess_mime(filename, content)
    if mime == "application/pdf":
        return "pdf"
    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return "docx"
    if mime in SUPPORTED_IMAGE_MIMES:
        return "image"
    return "unsupported"
